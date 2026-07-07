from __future__ import annotations
import json
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel

from . import deps
from .deps import BASE_DIR, get_engine, get_current_user_info
from modules.data.safe_io import atomic_write_json, load_json_safe
from modules.security import decrypt_config_keys
from modules.security import _encrypt_config_dict as _encrypt_config
from modules.narrative_style import NarrativeStyleManager, BUILTIN_STYLES

logger = logging.getLogger("chronoverse")
router = APIRouter(prefix="/api")

# 叙事风格管理器实例
_style_manager = NarrativeStyleManager(config_path=BASE_DIR / "config.json")


def _require_admin():
    """云版：敏感配置接口仅管理员可用"""
    info = get_current_user_info() or {}
    if not info.get("is_admin"):
        raise HTTPException(status_code=403, detail="仅管理员可操作配置")


def _write_config(config: dict):
    """写入 config.json 并通知 GameEngine 刷新缓存，同时清空 deps 缓存"""
    config_path = BASE_DIR / "config.json"
    atomic_write_json(config_path, config, indent=2, ensure_ascii=False)
    # [v9] Bug H4b: 同步清空 deps 中的配置缓存，避免读到旧值
    deps._config_cache = None
    eng = get_engine()
    if eng:
        eng.invalidate_config_cache()


def _mask_key(key: str) -> str:
    """[v9] Bug H7c: 对 API Key 做脱敏处理"""
    if not key or len(key) <= 8:
        return "***"
    return key[:4] + "***" + key[-4:]


def _mask_config_keys(config: dict) -> dict:
    """[v9] Bug H7c: 对配置中的所有 api_key 字段做脱敏。
    [v10.5] 新增 embedding 段的脱敏。
    [v10.5+] 新增 dialogue_llm 段的脱敏。"""
    masked = json.loads(json.dumps(config, ensure_ascii=False))
    for section in ("llm", "cheap_llm", "dialogue_llm", "image", "embedding"):
        if isinstance(masked.get(section), dict) and masked[section].get("api_key"):
            masked[section]["api_key"] = _mask_key(masked[section]["api_key"])
    for profile_type in ("llm_profiles", "image_profiles", "cheap_llm_profiles", "dialogue_llm_profiles"):
        profiles = masked.get(profile_type, {})
        if isinstance(profiles, dict):
            for name, profile in profiles.items():
                if isinstance(profile, dict) and profile.get("api_key"):
                    profile["api_key"] = _mask_key(profile["api_key"])
    return masked


class ConfigRequest(BaseModel):
    api_key: str
    base_url: str = "https://token-plan-cn.xiaomimimo.com/v1"
    model_name: str = "mimo-V2.5-Pro"


class ApplyProfileRequest(BaseModel):
    name: str
    target: str = "llm"


class SaveProfileRequest(BaseModel):
    name: str
    target: str = "llm"
    api_key: str = ""
    base_url: str = ""
    model_name: str = ""


class DeleteProfileRequest(BaseModel):
    name: str
    target: str = "llm"


class FullSettingsRequest(BaseModel):
    llm_api_key: str = ""
    llm_base_url: str = "https://token-plan-cn.xiaomimimo.com/v1"
    llm_model: str = "mimo-V2.5-Pro"
    llm_max_tokens: int = 0  # 0 = 不限制，使用API默认值
    cheap_llm_enabled: bool = False
    cheap_llm_api_key: str = ""
    cheap_llm_base_url: str = ""
    cheap_llm_model: str = ""
    # [v10.5+] 对话模型：用于游戏内叙事/NPC对话；未启用时回退到主力模型
    dialogue_llm_enabled: bool = False
    dialogue_llm_api_key: str = ""
    dialogue_llm_base_url: str = ""
    dialogue_llm_model: str = ""
    image_api_key: str = ""
    image_base_url: str = "https://api.siliconflow.cn/v1/images/generations"
    image_model: str = "Kwai-Kolors/Kolors"
    # [v10.5] 文本向量嵌入模型配置
    embedding_api_key: str = ""
    embedding_base_url: str = "https://api.siliconflow.cn/v1"
    embedding_model: str = "BAAI/bge-m3"
    theme: str = "dark"
    accent_color: str = "#c9a96e"
    bg_color: str = "#0a0a0f"
    text_color: str = "#e0d5c1"
    panel_bg: str = "#111120"
    font_size: str = "medium"
    fixed_prompt: str = ""
    fixed_prompt_enabled: bool = True
    strip_gray_narrative: bool = True
    npc_info_visibility: str = "immersive"
    economy_enabled: bool = False
    narrative_max_chars: int = 1000  # [v10.6] 叙事最大字数
    streaming_enabled: bool = True  # [v11] 流式输出开关
    action_validation_enabled: bool = True  # [v11] 行动合理性校验开关
    v10: dict = {}  # [v10.6+] v10高级配置


@router.get("/config")
async def load_config():
    config_path = BASE_DIR / "config.json"
    if not config_path.exists():
        return {}
    try:
        config = json.loads(config_path.read_text(encoding="utf-8"))
        config = decrypt_config_keys(config)
    except Exception as e:
        logger.error("Failed to load config: %s", e)
        return {}
    active_llm = config.get("active_llm_profile", "")
    active_image = config.get("active_image_profile", "")
    active_cheap = config.get("active_cheap_llm_profile", "")
    active_dialogue = config.get("active_dialogue_llm_profile", "")
    if active_llm and active_llm in config.get("llm_profiles", {}):
        profile = config["llm_profiles"][active_llm]
        config.setdefault("llm", {})
        config["llm"]["api_key"] = profile.get("api_key", config["llm"].get("api_key", ""))
        config["llm"]["base_url"] = profile.get("base_url", config["llm"].get("base_url", ""))
        config["llm"]["model_name"] = profile.get("model_name", config["llm"].get("model_name", ""))
    if active_cheap and active_cheap in config.get("cheap_llm_profiles", {}):
        profile = config["cheap_llm_profiles"][active_cheap]
        config.setdefault("cheap_llm", {})
        config["cheap_llm"]["api_key"] = profile.get("api_key", config["cheap_llm"].get("api_key", ""))
        config["cheap_llm"]["base_url"] = profile.get("base_url", config["cheap_llm"].get("base_url", ""))
        config["cheap_llm"]["model_name"] = profile.get("model_name", config["cheap_llm"].get("model_name", ""))
    if active_dialogue and active_dialogue in config.get("dialogue_llm_profiles", {}):
        profile = config["dialogue_llm_profiles"][active_dialogue]
        config.setdefault("dialogue_llm", {})
        config["dialogue_llm"]["api_key"] = profile.get("api_key", config["dialogue_llm"].get("api_key", ""))
        config["dialogue_llm"]["base_url"] = profile.get("base_url", config["dialogue_llm"].get("base_url", ""))
        config["dialogue_llm"]["model_name"] = profile.get("model_name", config["dialogue_llm"].get("model_name", ""))
    if active_image and active_image in config.get("image_profiles", {}):
        profile = config["image_profiles"][active_image]
        config.setdefault("image", {})
        config["image"]["api_key"] = profile.get("api_key", config["image"].get("api_key", ""))
        config["image"]["base_url"] = profile.get("base_url", config["image"].get("base_url", ""))
        config["image"]["model_name"] = profile.get("model_name", config["image"].get("model_name", ""))
    # [v9] Bug H7c: 返回前对所有 api_key 字段脱敏
    return _mask_config_keys(config)


@router.get("/config/raw")
async def load_config_raw():
    """获取未脱敏的配置（仅管理员可用，云版保护API Key）"""
    _require_admin()
    config_path = BASE_DIR / "config.json"
    if not config_path.exists():
        return {}
    try:
        config = json.loads(config_path.read_text(encoding="utf-8"))
        config = decrypt_config_keys(config)
    except Exception as e:
        logger.error("Failed to load config: %s", e)
        return {}
    active_llm = config.get("active_llm_profile", "")
    active_cheap = config.get("active_cheap_llm_profile", "")
    active_dialogue = config.get("active_dialogue_llm_profile", "")
    if active_llm and active_llm in config.get("llm_profiles", {}):
        profile = config["llm_profiles"][active_llm]
        config.setdefault("llm", {})
        config["llm"]["api_key"] = profile.get("api_key", config["llm"].get("api_key", ""))
        config["llm"]["base_url"] = profile.get("base_url", config["llm"].get("base_url", ""))
        config["llm"]["model_name"] = profile.get("model_name", config["llm"].get("model_name", ""))
    if active_cheap and active_cheap in config.get("cheap_llm_profiles", {}):
        profile = config["cheap_llm_profiles"][active_cheap]
        config.setdefault("cheap_llm", {})
        config["cheap_llm"]["api_key"] = profile.get("api_key", config["cheap_llm"].get("api_key", ""))
        config["cheap_llm"]["base_url"] = profile.get("base_url", config["cheap_llm"].get("base_url", ""))
        config["cheap_llm"]["model_name"] = profile.get("model_name", config["cheap_llm"].get("model_name", ""))
    if active_dialogue and active_dialogue in config.get("dialogue_llm_profiles", {}):
        profile = config["dialogue_llm_profiles"][active_dialogue]
        config.setdefault("dialogue_llm", {})
        config["dialogue_llm"]["api_key"] = profile.get("api_key", config["dialogue_llm"].get("api_key", ""))
        config["dialogue_llm"]["base_url"] = profile.get("base_url", config["dialogue_llm"].get("base_url", ""))
        config["dialogue_llm"]["model_name"] = profile.get("model_name", config["dialogue_llm"].get("model_name", ""))
    return config


@router.post("/config")
async def save_config(req: ConfigRequest):
    _require_admin()
    config_path = BASE_DIR / "config.json"
    config = {}
    if config_path.exists():
        try:
            config = json.loads(config_path.read_text(encoding="utf-8"))
            config = decrypt_config_keys(config)
        except Exception as e:
            logger.warning("Failed to load existing config: %s", e)
            config = {}
    config.setdefault("llm", {})
    config["llm"]["api_key"] = req.api_key
    config["llm"]["base_url"] = req.base_url
    config["llm"]["model_name"] = req.model_name
    try:
        _encrypt_config(config)
        _write_config(config)
    except Exception as e:
        logger.error("Failed to save config: %s", e)
        return {"error": f"保存配置失败: {e}"}
    return {"status": "ok"}


@router.get("/model-profiles")
async def get_model_profiles():
    config_path = BASE_DIR / "config.json"
    if not config_path.exists():
        return {"llm_profiles": {}, "image_profiles": {}, "cheap_llm_profiles": {}, "dialogue_llm_profiles": {}, "active_llm": "", "active_image": "", "active_cheap_llm": "", "active_dialogue_llm": ""}
    config = json.loads(config_path.read_text(encoding="utf-8"))
    # [v9] Bug H7c: 对 profile 中的 api_key 脱敏
    result = {
        "llm_profiles": config.get("llm_profiles", {}),
        "image_profiles": config.get("image_profiles", {}),
        "cheap_llm_profiles": config.get("cheap_llm_profiles", {}),
        "dialogue_llm_profiles": config.get("dialogue_llm_profiles", {}),
        "active_llm": config.get("active_llm_profile", ""),
        "active_image": config.get("active_image_profile", ""),
        "active_cheap_llm": config.get("active_cheap_llm_profile", ""),
        "active_dialogue_llm": config.get("active_dialogue_llm_profile", ""),
    }
    for profile_type in ("llm_profiles", "image_profiles", "cheap_llm_profiles", "dialogue_llm_profiles"):
        profiles = result.get(profile_type, {})
        if isinstance(profiles, dict):
            for name, profile in profiles.items():
                if isinstance(profile, dict) and profile.get("api_key"):
                    profile["api_key"] = _mask_key(profile["api_key"])
    return result


@router.post("/model-profiles/apply")
async def apply_model_profile(req: ApplyProfileRequest):
    _require_admin()
    config_path = BASE_DIR / "config.json"
    if not config_path.exists():
        return {"error": "配置文件不存在"}
    config = json.loads(config_path.read_text(encoding="utf-8"))
    # [v9] Bug H20: 读取后先解密，避免已加密的 api_key 被二次加密
    config = decrypt_config_keys(config)
    if req.target == "llm":
        profiles = config.get("llm_profiles", {})
        if req.name not in profiles:
            return {"error": "配置不存在"}
        config.setdefault("llm", {})
        for k, v in profiles[req.name].items():
            config["llm"][k] = v
        config["active_llm_profile"] = req.name
    elif req.target == "cheap":
        profiles = config.get("cheap_llm_profiles", {})
        if req.name not in profiles:
            return {"error": "配置不存在"}
        config.setdefault("cheap_llm", {})
        for k, v in profiles[req.name].items():
            config["cheap_llm"][k] = v
        config["active_cheap_llm_profile"] = req.name
    elif req.target == "dialogue":
        profiles = config.get("dialogue_llm_profiles", {})
        if req.name not in profiles:
            return {"error": "配置不存在"}
        config.setdefault("dialogue_llm", {})
        for k, v in profiles[req.name].items():
            config["dialogue_llm"][k] = v
        config["active_dialogue_llm_profile"] = req.name
    elif req.target == "image":
        profiles = config.get("image_profiles", {})
        if req.name not in profiles:
            return {"error": "配置不存在"}
        config.setdefault("image", {})
        for k, v in profiles[req.name].items():
            config["image"][k] = v
        config["active_image_profile"] = req.name
    else:
        return {"error": "未知类型"}
    _encrypt_config(config)
    _write_config(config)
    # [v10.6] 热更新 LLM 配置：profile 切换后重建 LLM 实例
    eng = get_engine()
    if eng and hasattr(eng, 'reload_llm_from_config'):
        try:
            eng.reload_llm_from_config()
        except Exception as e:
            logger.error("Failed to hot-reload LLM after profile apply: %s", e)
    return {"status": "ok", "llm": config.get("llm", {}), "image": config.get("image", {}), "cheap_llm": config.get("cheap_llm", {}), "dialogue_llm": config.get("dialogue_llm", {})}


@router.post("/model-profiles/save")
async def save_model_profile(req: SaveProfileRequest):
    _require_admin()
    config_path = BASE_DIR / "config.json"
    if not config_path.exists():
        return {"error": "配置文件不存在"}
    config = json.loads(config_path.read_text(encoding="utf-8"))
    # [v9] Bug H20: 读取后先解密，避免已加密的 api_key 被二次加密
    config = decrypt_config_keys(config)
    if req.target == "llm":
        config.setdefault("llm_profiles", {})
        config["llm_profiles"][req.name] = {
            "api_key": req.api_key,
            "base_url": req.base_url,
            "model_name": req.model_name,
        }
        config["active_llm_profile"] = req.name
        config.setdefault("llm", {})
        config["llm"]["api_key"] = req.api_key
        config["llm"]["base_url"] = req.base_url
        config["llm"]["model_name"] = req.model_name
    elif req.target == "cheap":
        config.setdefault("cheap_llm_profiles", {})
        config["cheap_llm_profiles"][req.name] = {
            "api_key": req.api_key,
            "base_url": req.base_url,
            "model_name": req.model_name,
        }
        config["active_cheap_llm_profile"] = req.name
        config.setdefault("cheap_llm", {})
        config["cheap_llm"]["api_key"] = req.api_key
        config["cheap_llm"]["base_url"] = req.base_url
        config["cheap_llm"]["model_name"] = req.model_name
    elif req.target == "dialogue":
        config.setdefault("dialogue_llm_profiles", {})
        config["dialogue_llm_profiles"][req.name] = {
            "api_key": req.api_key,
            "base_url": req.base_url,
            "model_name": req.model_name,
        }
        config["active_dialogue_llm_profile"] = req.name
        config.setdefault("dialogue_llm", {})
        config["dialogue_llm"]["api_key"] = req.api_key
        config["dialogue_llm"]["base_url"] = req.base_url
        config["dialogue_llm"]["model_name"] = req.model_name
    elif req.target == "image":
        config.setdefault("image_profiles", {})
        config["image_profiles"][req.name] = {
            "api_key": req.api_key,
            "base_url": req.base_url,
            "model_name": req.model_name,
        }
        config["active_image_profile"] = req.name
        config.setdefault("image", {})
        config["image"]["api_key"] = req.api_key
        config["image"]["base_url"] = req.base_url
        config["image"]["model_name"] = req.model_name
    _encrypt_config(config)
    _write_config(config)
    # [v10.6] 热更新 LLM 配置：保存 profile 后重建 LLM 实例
    eng = get_engine()
    if eng and hasattr(eng, 'reload_llm_from_config'):
        try:
            eng.reload_llm_from_config()
        except Exception as e:
            logger.error("Failed to hot-reload LLM after profile save: %s", e)
    return {"status": "ok"}


@router.post("/model-profiles/delete")
async def delete_model_profile(req: DeleteProfileRequest):
    _require_admin()
    config_path = BASE_DIR / "config.json"
    if not config_path.exists():
        return {"error": "配置文件不存在"}
    config = json.loads(config_path.read_text(encoding="utf-8"))
    # [v9] Bug H20: 读取后先解密，避免已加密的 api_key 被二次加密
    config = decrypt_config_keys(config)
    if req.target == "llm":
        profiles = config.get("llm_profiles", {})
        if req.name in profiles:
            del profiles[req.name]
            config["llm_profiles"] = profiles
        if config.get("active_llm_profile") == req.name:
            config["active_llm_profile"] = ""
    elif req.target == "cheap":
        profiles = config.get("cheap_llm_profiles", {})
        if req.name in profiles:
            del profiles[req.name]
            config["cheap_llm_profiles"] = profiles
        if config.get("active_cheap_llm_profile") == req.name:
            config["active_cheap_llm_profile"] = ""
    elif req.target == "dialogue":
        profiles = config.get("dialogue_llm_profiles", {})
        if req.name in profiles:
            del profiles[req.name]
            config["dialogue_llm_profiles"] = profiles
        if config.get("active_dialogue_llm_profile") == req.name:
            config["active_dialogue_llm_profile"] = ""
    elif req.target == "image":
        profiles = config.get("image_profiles", {})
        if req.name in profiles:
            del profiles[req.name]
            config["image_profiles"] = profiles
        if config.get("active_image_profile") == req.name:
            config["active_image_profile"] = ""
    _encrypt_config(config)
    _write_config(config)
    return {"status": "ok"}


@router.post("/settings")
async def save_settings(req: ConfigRequest):
    _require_admin()
    config_path = BASE_DIR / "config.json"
    config = {}
    if config_path.exists():
        try:
            config = json.loads(config_path.read_text(encoding="utf-8"))
        except Exception:
            config = {}
    # [v9] Bug H20: 读取后先解密，避免已加密的 api_key 被二次加密
    config = decrypt_config_keys(config)
    config.setdefault("llm", {})
    config["llm"]["api_key"] = req.api_key
    config["llm"]["base_url"] = req.base_url
    config["llm"]["model_name"] = req.model_name
    _encrypt_config(config)
    _write_config(config)
    return {"status": "ok"}


@router.post("/full-settings")
async def save_full_settings(req: FullSettingsRequest):
    _require_admin()
    config_path = BASE_DIR / "config.json"
    config = {}
    if config_path.exists():
        try:
            config = json.loads(config_path.read_text(encoding="utf-8"))
        except Exception:
            config = {}
    # [v9] Bug H20: 读取后先解密，避免已加密的 api_key 被二次加密
    config = decrypt_config_keys(config)
    config.setdefault("llm", {})
    config["llm"]["api_key"] = req.llm_api_key
    config["llm"]["base_url"] = req.llm_base_url
    config["llm"]["model_name"] = req.llm_model
    config["llm"]["max_tokens"] = req.llm_max_tokens  # 0 = 不限制
    config.setdefault("cheap_llm", {})
    config["cheap_llm"]["enabled"] = req.cheap_llm_enabled
    config["cheap_llm"]["api_key"] = req.cheap_llm_api_key if req.cheap_llm_enabled else ""
    config["cheap_llm"]["base_url"] = req.cheap_llm_base_url if req.cheap_llm_enabled else ""
    config["cheap_llm"]["model_name"] = req.cheap_llm_model if req.cheap_llm_enabled else ""
    # [v10.5+] 保存对话模型配置
    config.setdefault("dialogue_llm", {})
    config["dialogue_llm"]["enabled"] = req.dialogue_llm_enabled
    config["dialogue_llm"]["api_key"] = req.dialogue_llm_api_key if req.dialogue_llm_enabled else ""
    config["dialogue_llm"]["base_url"] = req.dialogue_llm_base_url if req.dialogue_llm_enabled else ""
    config["dialogue_llm"]["model_name"] = req.dialogue_llm_model if req.dialogue_llm_enabled else ""
    config.setdefault("image", {})
    config["image"]["api_key"] = req.image_api_key
    config["image"]["base_url"] = req.image_base_url
    config["image"]["model_name"] = req.image_model
    # [v10.5] 保存文本向量嵌入模型配置
    config.setdefault("embedding", {})
    config["embedding"]["api_key"] = req.embedding_api_key
    config["embedding"]["base_url"] = req.embedding_base_url
    config["embedding"]["model_name"] = req.embedding_model
    config.setdefault("ui", {})
    config["ui"]["theme"] = req.theme
    config["ui"]["accent_color"] = req.accent_color
    config["ui"]["bg_color"] = req.bg_color
    config["ui"]["text_color"] = req.text_color
    config["ui"]["panel_bg"] = req.panel_bg
    config["ui"]["font_size"] = req.font_size
    config.setdefault("fixed_prompt", {})
    config["fixed_prompt"]["content"] = req.fixed_prompt
    config["fixed_prompt"]["enabled"] = req.fixed_prompt_enabled
    config["ui"]["strip_gray_narrative"] = req.strip_gray_narrative
    config["npc_info_visibility"] = req.npc_info_visibility
    config.setdefault("game", {})
    config["game"]["economy_enabled"] = req.economy_enabled
    config["game"]["narrative_max_chars"] = req.narrative_max_chars  # [v10.6]
    config["game"]["streaming_enabled"] = req.streaming_enabled  # [v11]
    config["game"]["action_validation_enabled"] = req.action_validation_enabled  # [v11]
    # [v10.6+] 保存v10高级配置
    if req.v10:
        config["v10"] = req.v10
    _encrypt_config(config)
    _write_config(config)
    eng = get_engine()
    if eng and hasattr(eng, 'npc_registry') and eng.npc_registry:
        eng.npc_registry.set_info_visibility(req.npc_info_visibility)
    # [v10.6] 运行时更新叙事字数，设置修改后立即生效
    if eng and hasattr(eng, 'narrative_max_chars'):
        eng.narrative_max_chars = req.narrative_max_chars
    # [Bug] 运行时更新 max_tokens，设置修改后立即生效
    if hasattr(eng, 'main_llm') and eng.main_llm:
        eng.main_llm.set_default_max_tokens(req.llm_max_tokens)
    # [v10.6] 热更新 LLM 配置：重建 LLM 实例，使 API Key / Base URL / Model Name 变更即时生效
    # [v11] 加 _game_lock 防止热更新期间正在进行的 LLM 调用因连接池关闭而失败
    if eng and hasattr(eng, '_game_lock') and hasattr(eng, 'reload_llm_from_config'):
        async with eng._game_lock:
            try:
                eng.reload_llm_from_config()
            except Exception as e:
                logger.error("Failed to hot-reload LLM config: %s", e)
    # [v10.6+] 运行时应用v10高级配置
    if eng and req.v10:
        try:
            _apply_v10_config(eng, req.v10)
        except Exception as e:
            logger.error("Failed to apply v10 config: %s", e)
    return {"status": "ok"}


def _apply_v10_config(eng, v10_config: dict):
    """将v10高级配置应用到引擎运行时"""
    # 叙事审查器
    if 'narrative_reviewer' in v10_config:
        rv = v10_config['narrative_reviewer']
        if hasattr(eng, 'narrative_reviewer'):
            eng.narrative_reviewer.enabled = rv.get('enabled', True)
            eng.narrative_reviewer.review_interval = rv.get('review_interval', 10)
            eng.narrative_reviewer.max_lessons = rv.get('max_lessons', 30)
    # NPC程序记忆
    if 'npc_procedural_memory' in v10_config:
        npc = v10_config['npc_procedural_memory']
        if hasattr(eng, 'npc_procedural_memory'):
            eng.npc_procedural_memory.enabled = npc.get('enabled', True)
            eng.npc_procedural_memory.max_entries_per_npc = npc.get('max_entries_per_npc', 30)
            eng.npc_procedural_memory.evolve_interval_days = npc.get('evolve_interval_days', 10)
    # 世界任务板
    if 'world_task_board' in v10_config:
        tb = v10_config['world_task_board']
        if hasattr(eng, 'world_task_board'):
            eng.world_task_board.enabled = tb.get('enabled', True)
            eng.world_task_board.max_active_tasks = tb.get('max_active_tasks', 20)
            eng.world_task_board.auto_assign = tb.get('auto_assign', True)
    # 记忆整理器
    if 'memory_curator' in v10_config:
        mc = v10_config['memory_curator']
        if hasattr(eng, 'memory_curator'):
            eng.memory_curator.enabled = mc.get('enabled', True)
            eng.memory_curator.curate_interval = mc.get('curate_interval', 15)
            eng.memory_curator.max_archived_memories = mc.get('max_archived_memories', 200)
    # 蝴蝶审批门
    if 'butterfly_approval_gate' in v10_config:
        bf = v10_config['butterfly_approval_gate']
        if hasattr(eng, 'butterfly_approval_gate'):
            eng.butterfly_approval_gate.enabled = bf.get('enabled', False)
            eng.butterfly_approval_gate.approval_threshold = bf.get('approval_threshold', 7.0)
    # 伏笔生命周期
    if 'foreshadow_lifecycle' in v10_config:
        fs = v10_config['foreshadow_lifecycle']
        if hasattr(eng, 'foreshadow_lifecycle'):
            eng.foreshadow_lifecycle.enabled = fs.get('enabled', True)
            eng.foreshadow_lifecycle.reminder_mode = fs.get('reminder_mode', 'normal')
            eng.foreshadow_lifecycle.stale_threshold_days = fs.get('stale_threshold_days', 30)
    # 连续性审计
    if 'continuity_auditor' in v10_config:
        ca = v10_config['continuity_auditor']
        if hasattr(eng, 'continuity_auditor'):
            eng.continuity_auditor.enabled = ca.get('enabled', True)
    # 多智能体叙事
    if 'multi_agent_narrative' in v10_config:
        ma = v10_config['multi_agent_narrative']
        if hasattr(eng, 'multi_agent_narrative'):
            eng.multi_agent_narrative.enabled = ma.get('enabled', True)
            eng.multi_agent_narrative.max_revisions = ma.get('max_revisions', 1)
        # 触发灵敏度：保存在引擎级变量，供 TurnProcessorV2 使用
        eng.multi_agent_sensitivity = ma.get('sensitivity', 'low')


@router.get("/stats")
async def get_engine_stats():
    """获取引擎运行统计：LLM调用、缓存命中、任务队列"""
    eng = get_engine()
    if not eng:
        return {"status": "no_engine"}
    result = {
        "task_queue": eng.task_queue.get_stats() if hasattr(eng, 'task_queue') and eng.task_queue else None,
        "llm": None,
        "main_model": None,
        "cheap_model": None,
        "dialogue_model": None,
        "embedding": None,
    }
    if eng.llm:
        result["llm"] = eng.llm.get_stats() if hasattr(eng.llm, 'get_stats') else {}
    if hasattr(eng, 'main_llm') and eng.main_llm:
        result["main_model"] = {
            "name": eng.main_llm.model_name,
            **eng.main_llm.get_stats()
        }
    if hasattr(eng, 'cheap_llm') and eng.cheap_llm:
        result["cheap_model"] = {
            "name": eng.cheap_llm.model_name,
            **eng.cheap_llm.get_stats()
        }
    if hasattr(eng, 'dialogue_llm') and eng.dialogue_llm:
        result["dialogue_model"] = {
            "name": eng.dialogue_llm.model_name,
            **eng.dialogue_llm.get_stats()
        }
    # [P3-7] 嵌入服务统计
    if hasattr(eng, '_embedding_function') and eng._embedding_function:
        ef = eng._embedding_function
        result["embedding"] = {
            "model": getattr(ef, 'model_name', 'unknown'),
            "cache_size": len(ef._cache) if hasattr(ef, '_cache') else 0,
            "api_calls": getattr(ef, '_api_call_count', 0),
            "cache_hits": getattr(ef, '_cache_hit_count', 0),
            "fallback_count": getattr(ef, '_fallback_count', 0),
        }
    return result


@router.get("/health")
async def health_check():
    """[P3-7] 健康检查端点：验证各服务可用性"""
    eng = get_engine()
    health = {
        "status": "healthy",
        "checks": {},
        "timestamp": __import__('datetime').datetime.now().isoformat(),
    }

    # 检查 engine 是否初始化
    if not eng:
        health["status"] = "degraded"
        health["checks"]["engine"] = {"status": "down", "error": "engine not initialized"}
        return health
    health["checks"]["engine"] = {"status": "up"}

    # 检查 LLM 是否可用
    if hasattr(eng, 'main_llm') and eng.main_llm:
        health["checks"]["llm"] = {"status": "up", "model": eng.main_llm.model_name}
    else:
        health["checks"]["llm"] = {"status": "down", "error": "main_llm not initialized"}
        health["status"] = "degraded"

    # [v10.5+] 检查对话模型（可选，未配置时不影响整体健康度）
    if hasattr(eng, 'dialogue_llm') and eng.dialogue_llm:
        health["checks"]["dialogue_llm"] = {"status": "up", "model": eng.dialogue_llm.model_name}
    else:
        health["checks"]["dialogue_llm"] = {"status": "skip", "note": "未配置，回退到主力模型"}

    # 检查嵌入服务
    if hasattr(eng, '_embedding_function') and eng._embedding_function:
        health["checks"]["embedding"] = {
            "status": "up",
            "model": getattr(eng._embedding_function, 'model_name', 'unknown'),
        }
    else:
        health["checks"]["embedding"] = {"status": "down", "error": "embedding not configured"}

    # 检查记忆存储
    if hasattr(eng, 'memory') and eng.memory:
        try:
            count = eng.memory.collection.count()
            health["checks"]["memory"] = {"status": "up", "count": count}
        except Exception as e:
            health["checks"]["memory"] = {"status": "down", "error": str(e)}
            health["status"] = "degraded"
    else:
        health["checks"]["memory"] = {"status": "down", "error": "memory not initialized"}

    # 检查任务队列
    if hasattr(eng, 'task_queue') and eng.task_queue:
        try:
            stats = eng.task_queue.get_stats()
            health["checks"]["task_queue"] = {"status": "up", **stats}
        except Exception as e:
            health["checks"]["task_queue"] = {"status": "down", "error": str(e)}
    else:
        health["checks"]["task_queue"] = {"status": "down", "error": "task_queue not initialized"}

    return health


@router.post("/upload-description")
async def upload_description(file: UploadFile = File(...)):
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    content_bytes = await file.read()
    text = ""
    if ext == "txt":
        text = content_bytes.decode("utf-8", errors="ignore")
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return {"error": f"Word文件解析失败: {e}"}
    else:
        return {"error": f"不支持的文件格式: .{ext}，仅支持 .txt 和 .docx"}
    return {"text": text, "filename": file.filename}


# ═══════════════════════════════════════════════════════════════
# [v9] 叙事风格管理 API
# ═══════════════════════════════════════════════════════════════

class SetStyleRequest(BaseModel):
    style_name: str
    custom_text: str = ""
    narrative_perspective: str = "third"  # second/third/first


class AddCustomStyleRequest(BaseModel):
    name: str
    description: str


class ExtractStyleRequest(BaseModel):
    text: str


@router.get("/narrative-styles")
async def get_narrative_styles():
    """获取所有可用叙事风格（预置 + 自定义）及当前激活风格"""
    _style_manager.invalidate_cache()
    config_path = BASE_DIR / "config.json"
    config = {}
    if config_path.exists():
        try:
            config = json.loads(config_path.read_text(encoding="utf-8"))
        except Exception:
            pass

    styles = {}
    # 预置风格
    for name, desc in BUILTIN_STYLES.items():
        styles[name] = {"description": desc, "builtin": True}
    # 自定义风格
    for name, desc in config.get("narrative_styles", {}).items():
        styles[name] = {"description": desc, "builtin": False}

    active = config.get("game", {}).get("narrative_style", "章回体")
    custom_text = config.get("game", {}).get("narrative_style_custom", "")

    return {
        "styles": styles,
        "active": active,
        "custom_text": custom_text,
    }


@router.post("/narrative-style")
async def set_narrative_style(req: SetStyleRequest):
    """设置当前叙事风格"""
    _style_manager.invalidate_cache()
    config_path = BASE_DIR / "config.json"
    config = {}
    if config_path.exists():
        try:
            config = json.loads(config_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    config.setdefault("game", {})
    config["game"]["narrative_style"] = req.style_name
    config["game"]["narrative_perspective"] = req.narrative_perspective
    if req.custom_text:
        config["game"]["narrative_style_custom"] = req.custom_text
    elif req.style_name != "自定义":
        config["game"]["narrative_style_custom"] = ""
    _write_config(config)
    return {"status": "ok", "active": req.style_name}


@router.post("/narrative-style/custom")
async def add_custom_style(req: AddCustomStyleRequest):
    """添加自定义叙事风格"""
    if req.name in BUILTIN_STYLES:
        return {"error": f"不能覆盖预置风格「{req.name}」"}
    if not req.name.strip():
        return {"error": "风格名称不能为空"}
    _style_manager.invalidate_cache()
    _style_manager.add_custom_style(req.name.strip(), req.description)
    return {"status": "ok", "name": req.name}


@router.delete("/narrative-style/custom/{name}")
async def delete_custom_style(name: str):
    """删除自定义叙事风格"""
    _style_manager.invalidate_cache()
    if name in BUILTIN_STYLES:
        return {"error": "不能删除预置风格"}
    ok = _style_manager.delete_custom_style(name)
    if not ok:
        return {"error": f"风格「{name}」不存在"}
    return {"status": "ok"}


@router.post("/narrative-style/extract")
async def extract_style_from_text(req: ExtractStyleRequest):
    """从用户提供的文本中提取写作风格特征"""
    engine = get_engine()
    llm = engine.llm if engine else None
    keywords = _style_manager.extract_style_keywords(req.text, llm=llm)
    return {"keywords": keywords}


@router.post("/narrative-style/upload")
async def upload_style_file(file: UploadFile = File(...)):
    """上传txt文件作为自定义写作风格要求"""
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    content_bytes = await file.read()
    text = ""
    if ext == "txt":
        text = content_bytes.decode("utf-8", errors="ignore")
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return {"error": f"Word文件解析失败: {e}"}
    else:
        return {"error": f"不支持的文件格式: .{ext}，仅支持 .txt 和 .docx"}

    if not text.strip():
        return {"error": "文件内容为空"}

    # 提取风格特征
    engine = get_engine()
    llm = engine.llm if engine else None
    keywords = _style_manager.extract_style_keywords(text, llm=llm)

    return {
        "text": text,
        "filename": file.filename,
        "extracted_style": keywords,
    }
